"""Extract text from CV files (PDF and DOCX).

DOCX extraction reads body paragraphs and tables together, since a CV's work
history is often laid out as a table; PDF extraction reads the text layer only
(scanned/image-only PDFs yield empty text and surface as `unreadable` upstream).
"""

import io
import re
import zipfile

from docx import Document
from pypdf import PdfReader

from app.services.archive import BoundedArchiveTooLarge, bounded_archive

# A multiple, not a hardcoded byte count: the caller's max_chars already
# encodes the app's configured limit, and we only need enough headroom to
# admit legitimately verbose documents (rich XML markup around a modest
# amount of text) while still rejecting a zip that inflates without bound.
# No `settings` import here — this module stays pure and the actual limit
# always comes from the caller.
_DOCX_BOMB_MULTIPLIER = 100

# A floor under the multiplier above. A stock DOCX's XML payload (styles,
# theme, fonts) runs to roughly 800 KB of fixed overhead regardless of how
# little body text it holds, so coupling the budget purely to
# max_chars * _DOCX_BOMB_MULTIPLIER means any configured max_chars below
# ~8,000 rejects every legitimate Word document, not just bombs. Do NOT
# delete this floor to "simplify" the budget expression — that silently
# reintroduces the bug where a small, real-world text limit makes every
# DOCX unreadable. 5 MB is generous for any real CV and still nowhere near
# an availability risk to the shared worker.
_DOCX_MIN_BUDGET_BYTES = 5 * 1024 * 1024


class UnsupportedDocument(Exception):
    """Raised when a file cannot be parsed despite being of a claimed kind."""

    pass


def sniff(data: bytes) -> str | None:
    """Detect file type from its contents.

    Returns "pdf", "docx", or None. Never trusts the filename or
    Content-Type header — only examines the bytes themselves.
    """
    # PDF starts with %PDF-
    if data.startswith(b"%PDF-"):
        return "pdf"

    # DOCX is a zip, but we don't accept any zip — only those containing
    # word/document.xml, to reject archives of unrelated content.
    if data.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as z:
                if "word/document.xml" in z.namelist():
                    return "docx"
        except (zipfile.BadZipFile, RuntimeError):
            # Not a valid zip or other zip errors; treat as unknown
            pass

    return None


def extract_text(data: bytes, kind: str, *, max_chars: int) -> str:
    """Extract text from a file of the given kind.

    A PDF or DOCX with no extractable text (e.g. a scan) returns "".
    Corrupt files raise UnsupportedDocument rather than leaking library errors.

    `max_chars` is not a formatting nicety — it bounds the work this function
    is allowed to do. This runs in a shared arq worker, so a hostile or
    malformed upload that forces unbounded decompression (a small DOCX whose
    document.xml inflates to hundreds of MB) or unbounded page iteration (a
    PDF with tens of thousands of pages) is a memory/availability risk to
    every tenant on that worker, not just the uploader. The caller supplies
    the bound from configuration; this module has no `settings` import and
    no opinion on what the number should be — it only enforces whatever it
    is given.

    Args:
        data: File bytes.
        kind: One of "pdf" or "docx".
        max_chars: Maximum number of characters to return. Required so the
            caller's configured limit is always the one enforced.

    Returns:
        Extracted text, possibly empty, never longer than max_chars.

    Raises:
        UnsupportedDocument: If the file cannot be parsed, or (for DOCX) if
            it actually decompresses past the bound derived from max_chars.
            Note "actually": the archive's own claims about its size are
            never consulted, because the attacker writes them.
    """
    if kind == "pdf":
        return _extract_pdf(data, max_chars=max_chars)
    elif kind == "docx":
        return _extract_docx(data, max_chars=max_chars)
    else:
        raise ValueError(f"Unsupported kind: {kind}")


def _extract_pdf(data: bytes, *, max_chars: int) -> str:
    """Extract text from a PDF, stopping once max_chars is reached.

    We accumulate text page by page and stop as soon as we have enough,
    rather than reading every page and slicing afterwards — a 50,000-page
    PDF must not force us to render every page before we notice we only
    needed the first few thousand characters.
    """
    try:
        reader = PdfReader(io.BytesIO(data))
        text_parts = []
        total_len = 0
        for page in reader.pages:
            if total_len >= max_chars:
                break
            text = page.extract_text()
            if text:
                text = _despace_letterspaced(text)
                text_parts.append(text)
                total_len += len(text)
        return "".join(text_parts)[:max_chars]
    except MemoryError:
        # Deliberately not caught by the bare `except Exception` below:
        # a corrupt PDF and an exhausted worker are both "we can't read
        # this", but we don't want a broad except to relabel an OOM as a
        # routine parse failure and mask the real cause.
        raise
    except Exception as e:
        # Bare Exception is a deliberate boundary here: pypdf can raise a
        # wide variety of internal error types for malformed input, and we
        # want all of them to surface as UnsupportedDocument rather than
        # leaking library internals to callers.
        raise UnsupportedDocument(f"Failed to parse PDF: {e}") from e


# A token of exactly one character (a letter or digit). Used to detect the
# letter-spaced layout some PDF generators (Canva, certain Google Docs
# exports) produce: every character is drawn at its own x-position, and
# pypdf then extracts "A d m i n i s t r a t i v e" instead of "Administrative".
_SINGLE_CHAR_TOKEN = re.compile(r"^[A-Za-z0-9]$")


def _despace_letterspaced(text: str) -> str:
    """Collapse single-character spacing produced by a PDF layout artifact.

    Some PDF generators draw each character at its own position, so the text
    layer reads "A d m i n i s t r a t i v e  2 0 1 8" rather than
    "Administrative 2018". Downstream every quote and every figure fails
    verification: the model quotes "Administrative" but the source only has
    "A d m i n i s t r a t i v e", and the amount check sees "2 0 1 8" as four
    separate digits. This is a rendering artifact, not real content, so it is
    repaired here at extraction time — the one place both the model's input
    and the verification source agree.

    Heuristic: a line whose tokens are predominantly single characters is
    letter-spaced. We only collapse within such lines, so normal prose (whose
    tokens are words) is untouched. Uppercase markers ("W O R K") are also
    letter-spaced and collapse to "WORK", which is what the author meant.
    """
    lines = text.split("\n")
    repaired = []
    for line in lines:
        # Letter-spaced lines use single spaces between letters and DOUBLE
        # spaces between words ("I  c o n s i d e r  m y s e l f  a").
        # Splitting on the single space would lose word boundaries; splitting
        # on the double space first keeps them, then each word's inner single
        # spaces are collapsed.
        tokens = [t for t in line.split("  ") if t]
        if len(tokens) < 2:
            # No word separators: maybe a normal prose line. Only collapse if
            # it is itself predominantly single characters.
            singles = sum(1 for t in line.split(" ") if _SINGLE_CHAR_TOKEN.match(t))
            total = len([t for t in line.split(" ") if t])
            if total >= 3 and singles / total >= 0.7:
                repaired.append("".join(t for t in line.split(" ") if t))
            else:
                repaired.append(line)
            continue
        # Count single-char tokens across the whole line to decide if this
        # line is letter-spaced at all.
        flat = line.split(" ")
        singles = sum(1 for t in flat if _SINGLE_CHAR_TOKEN.match(t))
        total = len([t for t in flat if t])
        if total >= 3 and singles / total >= 0.7:
            # Collapse each word's letter-spacing, keeping the word gaps.
            words = ["".join(t for t in word.split(" ") if t) for word in tokens]
            repaired.append(" ".join(words))
        else:
            repaired.append(line)
    return "\n".join(repaired)


def _extract_docx(data: bytes, *, max_chars: int) -> str:
    """Extract text from a DOCX file, refusing to decompress a bomb.

    The budget is derived from the caller's max_chars — see
    `_DOCX_BOMB_MULTIPLIER` — and is spent across the archive as a whole,
    so a thousand small members cannot do what one large member is
    forbidden to do.

    The actual bounded inflation lives in `app.services.archive`, shared
    with every other zip-based format we read (XLSX included). This
    function's only job on top of that is translating the archive module's
    format-neutral `BoundedArchiveTooLarge` into the `UnsupportedDocument`
    that every caller of this module already expects.

    Tables are read alongside paragraphs: many real CVs lay out work history
    or skills in a table rather than in body prose, and `python-docx` exposes
    `doc.paragraphs` and `doc.tables` as separate collections, so a loop over
    paragraphs alone would silently drop that content. Both collections share
    the single `max_chars` budget — a table-heavy CV is bounded the same way
    a prose one is.
    """
    try:
        budget = max(max_chars * _DOCX_BOMB_MULTIPLIER, _DOCX_MIN_BUDGET_BYTES)
        doc = Document(bounded_archive(data, budget=budget))
        text_parts: list[str] = []
        total_len = 0
        for para in doc.paragraphs:
            if total_len >= max_chars:
                break
            if para.text:
                text_parts.append(para.text)
                total_len += len(para.text) + 1  # +1 for the joining newline
        for table in doc.tables:
            if total_len >= max_chars:
                break
            for row in table.rows:
                if total_len >= max_chars:
                    break
                # A cell may itself contain paragraphs; take the cell's text
                # rather than recursing, so a heavily nested cell cannot run
                # away. Cells on a row are tab-separated to preserve the
                # column structure the LLM reads as employment history.
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not cells:
                    continue
                line = "\t".join(cells)
                text_parts.append(line)
                total_len += len(line) + 1
        return "\n".join(text_parts)[:max_chars]
    except BoundedArchiveTooLarge as e:
        raise UnsupportedDocument(str(e)) from e
    except UnsupportedDocument:
        raise
    except MemoryError:
        # See _extract_pdf: an OOM is not a "bad file" and must not be
        # relabeled as UnsupportedDocument by the broad except below.
        raise
    except Exception as e:
        # Bare Exception is a deliberate boundary here: python-docx and
        # zipfile can raise a variety of internal error types for malformed
        # input, and we want all of them to surface as UnsupportedDocument
        # rather than leaking library internals to callers.
        raise UnsupportedDocument(f"Failed to parse DOCX: {e}") from e
