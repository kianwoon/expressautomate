"""Turning an uploaded file into text, and refusing what we cannot read."""

import io
import struct
import zipfile

import pytest
from pypdf import PdfWriter

from app.services.cv.text import UnsupportedDocument, extract_text, sniff


def test_a_pdf_is_recognised_by_its_bytes():
    assert sniff(b"%PDF-1.7\nrest of file") == "pdf"


def test_a_bare_zip_is_not_a_docx():
    """A DOCX is a zip, so `PK` alone proves nothing.

    Accepting any archive means a recruiter can hand us a zip of holiday
    photos and get an unreadable job instead of a straight refusal.
    """
    assert sniff(b"PK\x03\x04" + b"\x00" * 64) is None


def test_a_real_docx_is_recognised(tmp_path):
    """Built here rather than committed as a fixture binary: the point is the
    presence of `word/document.xml` inside the archive, and that is clearer
    written than checked in."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("word/document.xml", "<w:document/>")
    assert sniff(buf.getvalue()) == "docx"


def test_something_that_is_neither_is_refused():
    assert sniff(b"just some text") is None


def test_pdf_with_no_extractable_text_returns_empty_string():
    """A scan or image-only PDF has no text to extract, so we return empty."""
    buf = io.BytesIO()
    writer = PdfWriter()
    # Create a blank page with no text
    writer.add_blank_page(width=612, height=792)
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    assert sniff(pdf_bytes) == "pdf"
    assert extract_text(pdf_bytes, "pdf", max_chars=1000) == ""


def test_corrupt_pdf_raises_unsupported_document():
    """A malformed PDF that cannot be parsed raises UnsupportedDocument."""
    corrupt_pdf = b"%PDF-1.4\n" + b"corrupted content" * 100

    assert sniff(corrupt_pdf) == "pdf"
    with pytest.raises(UnsupportedDocument):
        extract_text(corrupt_pdf, "pdf", max_chars=1000)


def _forge_uncompressed_size(archive: bytes, member: str, claimed: int) -> bytes:
    """Rewrite a member's declared uncompressed size, in both places a zip
    records it, leaving the compressed stream untouched.

    This is what a malicious zip does: the size fields are metadata written
    by whoever produced the file, and nothing in the format ties them to
    what the DEFLATE stream really yields.
    """
    raw = bytearray(archive)
    name = member.encode()

    # Central directory records: name length at +28, name at +46, size at +24.
    pos = 0
    while (pos := raw.find(b"PK\x01\x02", pos)) >= 0:
        (name_len,) = struct.unpack_from("<H", raw, pos + 28)
        if bytes(raw[pos + 46 : pos + 46 + name_len]) == name:
            struct.pack_into("<I", raw, pos + 24, claimed)
        pos += 4

    # Local headers: name length at +26, name at +30, size at +22.
    pos = 0
    while (pos := raw.find(b"PK\x03\x04", pos)) >= 0:
        (name_len,) = struct.unpack_from("<H", raw, pos + 26)
        if bytes(raw[pos + 30 : pos + 30 + name_len]) == name:
            struct.pack_into("<I", raw, pos + 22, claimed)
        pos += 4

    return bytes(raw)


def test_docx_that_lies_about_its_size_is_still_refused():
    """The bypass that a metadata-only guard cannot see.

    This member *declares* 100 bytes, so summing `ZipFile.infolist()`
    file_size finds nothing wrong — but its DEFLATE stream really expands
    to 50MB, far past the bound implied by max_chars. The archive is only
    ~50KB on disk, so a guard that trusts the declaration would wave it
    through to python-docx and let it inflate unwatched.

    Built here with zipfile rather than committed as a binary: the point is
    the mismatch between claim and reality, and that is clearer in code.
    """
    from docx import Document as DocxDocument

    # Start from a genuinely valid document, so nothing *else* about this
    # file gives python-docx a reason to refuse it. Without this the test
    # would pass for the wrong reason — any old parse failure would satisfy
    # `pytest.raises` and we would learn nothing about the bound.
    doc = DocxDocument()
    doc.add_paragraph("A perfectly ordinary CV.")
    buf = io.BytesIO()
    doc.save(buf)

    with zipfile.ZipFile(buf, "a", zipfile.ZIP_DEFLATED) as z:
        z.writestr("word/media/bomb.bin", b"\x00" * (50 * 1024 * 1024))

    docx_bytes = _forge_uncompressed_size(buf.getvalue(), "word/media/bomb.bin", 100)

    # The declaration now looks entirely innocent to a metadata-only check.
    assert len(docx_bytes) < 1024 * 1024
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        assert sum(i.file_size for i in z.infolist()) < 20000 * 100

    assert sniff(docx_bytes) == "docx"
    with pytest.raises(UnsupportedDocument):
        extract_text(docx_bytes, "docx", max_chars=20000)


def test_docx_text_is_truncated_to_max_chars():
    """Extracted DOCX text never exceeds the caller's max_chars bound.

    Built with python-docx itself (rather than hand-rolled XML) so the
    relationships/content-types parts are all valid and the only thing
    under test is our truncation, not docx-format minutiae.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    long_paragraph = "word " * 500
    for _ in range(200):
        doc.add_paragraph(long_paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # A stock python-docx document already declares ~800KB of styles/theme
    # XML regardless of body content, so max_chars must be large enough
    # (as a real CV-text limit would be) that this legitimate overhead
    # doesn't itself trip the bomb check.
    assert sniff(docx_bytes) == "docx"
    full_length = len(long_paragraph) * 200
    result = extract_text(docx_bytes, "docx", max_chars=20000)
    assert len(result) <= 20000
    assert len(result) < full_length


def test_ordinary_docx_extracts_at_a_small_max_chars():
    """A stock python-docx file must not be rejected by a small text limit.

    A fresh document's fixed overhead (styles, theme, fonts) is roughly
    800KB of XML regardless of body content. If the decompression budget
    were purely `max_chars * _DOCX_BOMB_MULTIPLIER`, any max_chars below
    ~8,000 — a perfectly plausible configured limit — would make this
    ordinary, tiny CV unreadable. This is the case the floor exists for.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Recruitment Consultant, five years experience.")
    doc.add_paragraph("Contact: jane@example.com")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    assert sniff(docx_bytes) == "docx"
    result = extract_text(docx_bytes, "docx", max_chars=100)
    assert result != ""


def test_a_docx_table_is_extracted_not_dropped():
    """A CV that lays out its work history in a table must still yield text.

    `python-docx` exposes paragraphs and tables as separate collections, and a
    loop over paragraphs alone returns "" for a table-only document — which the
    pipeline would then mark `unreadable`. A table is a common CV layout, so
    this is the case that most directly loses a real candidate.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Priya Raman")  # a heading line in body prose
    table = doc.add_table(rows=3, cols=3)
    headers = ("Employer", "Title", "Years")
    for col, name in enumerate(headers):
        table.rows[0].cells[col].text = name
    rows = [
        ("Acme Logistics", "Operations Lead", "2019-2021"),
        ("Globex", "Warehouse Supervisor", "2021-Present"),
    ]
    for r, (employer, title, years) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = employer
        table.rows[r].cells[1].text = title
        table.rows[r].cells[2].text = years
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    assert sniff(docx_bytes) == "docx"
    result = extract_text(docx_bytes, "docx", max_chars=20000)
    # The body heading survives...
    assert "Priya Raman" in result
    # ...and so does every cell the table held — the whole point of the change.
    assert "Acme Logistics" in result
    assert "Operations Lead" in result
    assert "Globex" in result
    assert "Warehouse Supervisor" in result


def _pdf_with_text_pages(page_count: int, line: str) -> bytes:
    """Build a PDF whose every page carries real, extractable text.

    Written by hand rather than with a helper library because pypdf can
    only add *blank* pages, and a blank page is precisely what made the
    earlier version of the early-stop test vacuous: with no text anywhere,
    reading one page and reading fifty both produce "".
    """
    pages = []
    contents = []
    page_nums = []
    num = 3
    for _ in range(page_count):
        stream = f"BT /F1 12 Tf 72 720 Td ({line}) Tj ET".encode()
        content_num = num + 1
        pages.append(
            (
                num,
                b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                b"/Resources << /Font << /F1 %d 0 R >> >> /Contents %d 0 R >>"
                % (3 + page_count * 2, content_num),
            )
        )
        contents.append(
            (
                content_num,
                b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
            )
        )
        page_nums.append(b"%d 0 R" % num)
        num += 2

    font_num = 3 + page_count * 2
    objects = [
        (1, b"<< /Type /Catalog /Pages 2 0 R >>"),
        (
            2,
            b"<< /Type /Pages /Count %d /Kids [%s] >>"
            % (page_count, b" ".join(page_nums)),
        ),
        (font_num, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"),
        *pages,
        *contents,
    ]
    objects.sort()

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = {}
    for number, payload in objects:
        offsets[number] = out.tell()
        out.write(b"%d 0 obj\n" % number)
        out.write(payload)
        out.write(b"\nendobj\n")

    highest = max(offsets)
    xref = out.tell()
    out.write(b"xref\n0 %d\n" % (highest + 1))
    out.write(b"0000000000 65535 f \n")
    for i in range(1, highest + 1):
        if i in offsets:
            out.write(b"%010d 00000 n \n" % offsets[i])
        else:
            out.write(b"0000000000 65535 f \n")
    out.write(
        b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
        % (highest + 1, xref)
    )
    return out.getvalue()


def test_pdf_extraction_stops_early_rather_than_reading_every_page(monkeypatch):
    """A long PDF must stop being read once max_chars is satisfied.

    Every page here carries 80 real characters, so the fifty pages hold
    4000 — twenty times the 200 we ask for. Counting extract_text calls is
    the only way to see the difference: an implementation that reads all
    fifty pages and slices afterwards returns the same 200 characters as
    one that stops at page three, while doing twenty times the work on a
    worker shared with every other tenant. Assert on the work, not the
    output.
    """
    from pypdf._page import PageObject

    line = "x" * 80
    pdf_bytes = _pdf_with_text_pages(50, line)

    pages_read = 0
    real_extract = PageObject.extract_text

    def counting_extract(self, *args, **kwargs):
        nonlocal pages_read
        pages_read += 1
        return real_extract(self, *args, **kwargs)

    monkeypatch.setattr(PageObject, "extract_text", counting_extract)

    assert sniff(pdf_bytes) == "pdf"
    result = extract_text(pdf_bytes, "pdf", max_chars=200)

    assert len(result) == 200
    assert result.startswith(line)
    # Not pinned to the exact page count (brittle against pypdf internals) —
    # just proof that we stopped well short of reading all fifty pages.
    assert pages_read <= 10, f"read {pages_read} pages to satisfy 200 characters"


# --- Letter-spaced PDF text (Canva / Google Docs export artifact) ------------
# Some PDF generators draw every character at its own x-position; pypdf then
# extracts "A d m i n i s t r a t i v e  2 0 1 8" instead of
# "Administrative 2018". If that text reaches verification, every quote and
# every figure fails: the model quotes "Administrative" but the source has
# only the spaced form, and the amount check reads "2 0 1 8" as four digits.


def test_letterspaced_words_are_collapsed():
    from app.services.cv.text import _despace_letterspaced

    text = "A d m i n i s t r a t i v e  w o r k s"
    assert _despace_letterspaced(text) == "Administrative works"


def test_letterspaced_dates_are_collapsed():
    """The amount check reads '2 0 1 8' as four single digits; the collapsed
    '2018' is what the model quotes, so verification passes."""
    from app.services.cv.text import _despace_letterspaced

    text = "T E A M W O R K  G A R A G E  2 0 1 8  -  2 0 2 5"
    assert "2018" in _despace_letterspaced(text)
    assert "2025" in _despace_letterspaced(text)


def test_letterspaced_word_boundaries_are_preserved():
    """Word gaps in letter-spaced text are double spaces; collapsing must keep
    them, so 'I  c o n s i d e r' becomes 'I consider', not 'Iconsider'."""
    from app.services.cv.text import _despace_letterspaced

    text = "I  c o n s i d e r  m y s e l f  a"
    assert _despace_letterspaced(text) == "I consider myself a"


def test_normal_prose_is_untouched():
    """Ordinary text — words separated by single spaces — must not be altered."""
    from app.services.cv.text import _despace_letterspaced

    text = "Senior Operations Manager at Global Logistics"
    assert _despace_letterspaced(text) == text
