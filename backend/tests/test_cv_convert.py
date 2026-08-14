# allow-hardcode: the converted bytes and SQL below are test fixture content.
"""Legacy .doc → .docx conversion, wired into both parse paths.

No test here reaches LibreOffice: `maybe_convert` is monkeypatched to return a
canned .docx built with python-docx, so the suite never shells out to `soffice`.
The property that matters: a .doc (Word 97-2003) that `sniff` rejects is
rescued by conversion and reaches the parse, instead of going terminal
`unreadable` at the route or the job.
"""

import io
import uuid

import pytest
from sqlalchemy import text

from app.core.config import settings
from app.services.storage.r2 import InMemoryBodyStore
from app.workers import cv_jobs
from tests.conftest import AdminSessionLocal
from tests.test_candidate_roles_api import _a_candidate_row, agency  # noqa: F401

# The OLE2 Compound Document magic — a real .doc starts with this. Eight bytes
# is enough for `is_legacy_office` to say yes without a full file.
OLE2_DOC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64

# A minimal real .docx the extractor can read, built once and reused as the
# "converted" output the fake maybe_convert returns.


def _docx_bytes() -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Jane Tan")
    doc.add_paragraph("jane@example.com")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_DOCX = _docx_bytes()


def _fake_maybe_convert_that_returns(converted: bytes, kind: str = "docx"):
    async def _convert(data, *, kind):  # noqa: F811
        return converted, "docx"

    return _convert


def _fake_extraction(payload: dict):
    async def _extract(source, **kwargs):
        from app.services.cv.schema import CVResponse
        from app.services.llm.client import LLMResult

        return CVResponse.model_validate(payload), LLMResult(
            data=payload, model="test/fast", latency_ms=1, raw={}
        )

    return _extract


async def _seed(tenant_id, candidate_id, store, data: bytes):
    document_id = uuid.uuid4()
    key = f"{tenant_id}/documents/{document_id}.doc"
    if data is not None:
        await store.put_bytes(key, data, "application/msword")
    async with AdminSessionLocal() as s:
        await s.execute(
            text(
                "INSERT INTO candidate_documents (id, tenant_id, candidate_id, filename,"
                " content_type, byte_size, object_key, parse_state)"
                " VALUES (:i, :t, :c, 'cv.doc', 'application/msword', :b, :k, 'pending')"
            ),
            {"i": document_id, "t": tenant_id, "c": candidate_id, "b": len(data or b""), "k": key},
        )
        await s.commit()
    return document_id


async def _document(document_id):
    async with AdminSessionLocal() as s:
        return (
            await s.execute(
                text("SELECT parse_state, parse_error FROM candidate_documents WHERE id = :i"),
                {"i": document_id},
            )
        ).one()


async def _cleanup(tenant_id):
    async with AdminSessionLocal() as s:
        for table in ("extraction_evidence", "extractions", "candidate_documents"):
            await s.execute(text(f"DELETE FROM {table} WHERE tenant_id = :t"), {"t": tenant_id})
        await s.commit()


@pytest.fixture(autouse=True)
def _configured(monkeypatch):
    monkeypatch.setattr(settings, "LLM_PROVIDER_BASE_URL", "https://deepseek.test/v1")
    monkeypatch.setattr(settings, "LLM_PROVIDER_API_KEY", "test-key")
    monkeypatch.setattr(settings, "EXTRACTION_MODEL_FAST", "test/fast")


def _conversion_on(monkeypatch):
    """Turn conversion on without LibreOffice, patching the binary probe the gate
    reads through — `conversion_configured()` is a method, not a pydantic field,
    so setting it on the instance is rejected (see OCR's `ocr_on` fixture)."""
    from app.services.cv import convert as _convert_module

    monkeypatch.setattr(settings, "CV_CONVERT_ENABLED", True)
    monkeypatch.setattr(_convert_module, "converter_available", lambda: True)


def _conversion_off(monkeypatch):
    from app.services.cv import convert as _convert_module

    monkeypatch.setattr(settings, "CV_CONVERT_ENABLED", False)
    monkeypatch.setattr(_convert_module, "converter_available", lambda: True)


@pytest.fixture
def store(monkeypatch) -> InMemoryBodyStore:
    double = InMemoryBodyStore()
    monkeypatch.setattr(cv_jobs, "body_store", lambda: double)
    return double


# --- detection ---------------------------------------------------------------


def test_is_legacy_office_recognises_a_doc():
    from app.services.cv.convert import is_legacy_office

    assert is_legacy_office(OLE2_DOC) is True
    # A .docx is a zip and must NOT match — it is the modern format already.
    assert is_legacy_office(_DOCX) is False
    assert is_legacy_office(b"%PDF-1.4 not a doc") is False


# --- parse path --------------------------------------------------------------


@pytest.mark.asyncio
async def test_a_legacy_doc_is_converted_and_parses(
    agency, store, monkeypatch  # noqa: F811
):
    """A .doc that sniff rejects is rescued by conversion and reaches the parse."""
    tenant_id, user_id = agency
    candidate_id = await _a_candidate_row(tenant_id, user_id)
    document_id = await _seed(tenant_id, candidate_id, store, OLE2_DOC)

    _conversion_on(monkeypatch)
    monkeypatch.setattr(
        cv_jobs, "maybe_convert", _fake_maybe_convert_that_returns(_DOCX)
    )
    monkeypatch.setattr(
        cv_jobs, "extract_cv", _fake_extraction({"roles": [], "skills": []})
    )

    await cv_jobs.parse_candidate_cv(
        None, tenant_id=str(tenant_id), candidate_id=str(candidate_id),
        document_id=str(document_id),
    )

    row = await _document(document_id)
    # Converted to docx and parsed — not terminal unreadable.
    assert row.parse_state in ("parsed", "empty")
    await _cleanup(tenant_id)


@pytest.mark.asyncio
async def test_a_doc_with_conversion_disabled_is_unreadable_with_a_guide(
    agency, store, monkeypatch  # noqa: F811
):
    """Without LibreOffice, a .doc surfaces as an honest 'save as .docx' refusal."""
    tenant_id, user_id = agency
    candidate_id = await _a_candidate_row(tenant_id, user_id)
    document_id = await _seed(tenant_id, candidate_id, store, OLE2_DOC)

    _conversion_off(monkeypatch)

    async def _never(data, *, kind):  # noqa: F811
        raise AssertionError("conversion must not run when disabled")

    monkeypatch.setattr(cv_jobs, "maybe_convert", _never)

    await cv_jobs.parse_candidate_cv(
        None, tenant_id=str(tenant_id), candidate_id=str(candidate_id),
        document_id=str(document_id),
    )

    row = await _document(document_id)
    assert row.parse_state == "unreadable"
    assert "save it as" in row.parse_error.lower()
    await _cleanup(tenant_id)


@pytest.mark.asyncio
async def test_conversion_failure_is_unreadable_with_the_cause(
    agency, store, monkeypatch  # noqa: F811
):
    """A corrupt .doc LibreOffice cannot open surfaces the cause, not a crash."""
    from app.services.cv.convert import ConversionUnavailable

    tenant_id, user_id = agency
    candidate_id = await _a_candidate_row(tenant_id, user_id)
    document_id = await _seed(tenant_id, candidate_id, store, OLE2_DOC)

    _conversion_on(monkeypatch)

    async def _fails(data, *, kind):  # noqa: F811
        raise ConversionUnavailable("LibreOffice exited 1: corrupt input")

    monkeypatch.setattr(cv_jobs, "maybe_convert", _fails)

    await cv_jobs.parse_candidate_cv(
        None, tenant_id=str(tenant_id), candidate_id=str(candidate_id),
        document_id=str(document_id),
    )

    row = await _document(document_id)
    assert row.parse_state == "unreadable"
    assert "corrupt input" in row.parse_error
    await _cleanup(tenant_id)


@pytest.mark.asyncio
async def test_a_docx_is_not_touched_by_conversion(agency, store, monkeypatch):  # noqa: F811
    """A modern .docx skips conversion entirely — sniff already recognised it."""
    tenant_id, user_id = agency
    candidate_id = await _a_candidate_row(tenant_id, user_id)
    # Seed as a real .docx (key + content_type), so sniff recognises it and the
    # job never reaches the conversion branch.
    document_id = uuid.uuid4()
    key = f"{tenant_id}/documents/{document_id}.docx"
    await store.put_bytes(key, _DOCX, "application/docx")
    async with AdminSessionLocal() as s:
        await s.execute(
            text(
                "INSERT INTO candidate_documents (id, tenant_id, candidate_id, filename,"
                " content_type, byte_size, object_key, parse_state)"
                " VALUES (:i, :t, :c, 'cv.docx', 'application/docx', :b, :k, 'pending')"
            ),
            {"i": document_id, "t": tenant_id, "c": candidate_id, "b": len(_DOCX), "k": key},
        )
        await s.commit()

    async def _never(data, *, kind):  # noqa: F811
        raise AssertionError("a .docx must not reach conversion")

    monkeypatch.setattr(cv_jobs, "maybe_convert", _never)
    monkeypatch.setattr(
        cv_jobs, "extract_cv", _fake_extraction({"roles": [], "skills": []})
    )

    await cv_jobs.parse_candidate_cv(
        None, tenant_id=str(tenant_id), candidate_id=str(candidate_id),
        document_id=str(document_id),
    )

    row = await _document(document_id)
    assert row.parse_state in ("parsed", "empty")
    await _cleanup(tenant_id)
